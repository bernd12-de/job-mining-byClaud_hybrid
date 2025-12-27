"""
Multi-Format Document Parser v2.0
===================================

Unterstützt: PDF, DOCX, Google Docs (als DOCX)

Basierend auf:
- DOCX Skill Documentation (/mnt/skills/public/docx/SKILL.md)
- PyPDF2 für PDF-Verarbeitung
- python-docx für DOCX-Verarbeitung
- Pandoc für erweiterte DOCX-Konvertierung

Autor: Job Mining Project
Datum: 2025-11-01
"""

import re
import logging
from pathlib import Path
from typing import Optional, Dict, List
from dataclasses import dataclass
import subprocess
import tempfile

# Setup Logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Domain Model für geparste Dokumente"""
    filename: str
    file_type: str  # 'pdf', 'docx', 'gdoc'
    text: str
    word_count: int
    char_count: int
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        if self.word_count == 0:
            self.word_count = len(self.text.split())
        if self.char_count == 0:
            self.char_count = len(self.text)


class MultiFormatParser:
    """
    Universal Document Parser für Stellenanzeigen
    
    Unterstützte Formate:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Google Docs (als .docx export)
    
    Features:
    - Automatische Format-Erkennung
    - Text-Extraktion mit Struktur-Erhaltung
    - Metadaten-Extraktion
    - Fehlerbehandlung & Logging
    """
    
    def __init__(self):
        """Initialize Parser mit Format-Checks"""
        self._check_dependencies()
        logger.info("MultiFormatParser initialized")
    
    def _check_dependencies(self):
        """Prüft verfügbare Tools"""
        try:
            import PyPDF2
            self.pdf_available = True
        except ImportError:
            logger.warning("PyPDF2 not installed. Install: pip install PyPDF2 --break-system-packages")
            self.pdf_available = False
        
        try:
            import docx
            self.docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. Install: pip install python-docx --break-system-packages")
            self.docx_available = False
        
        # Check for pandoc (optional, für bessere DOCX-Konvertierung)
        try:
            subprocess.run(['pandoc', '--version'], 
                         capture_output=True, 
                         check=True,
                         timeout=5)
            self.pandoc_available = True
            logger.info("Pandoc detected - using enhanced DOCX conversion")
        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
            self.pandoc_available = False
            logger.info("Pandoc not available - using python-docx")
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """
        Haupt-Methode: Parst beliebiges Dokument
        
        Args:
            file_path: Pfad zur Datei (.pdf, .docx)
        
        Returns:
            ParsedDocument mit extrahiertem Text
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Format-Erkennung
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._parse_pdf(path)
        elif suffix in ['.docx', '.doc']:
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported format: {suffix}. Supported: .pdf, .docx")
    
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Extrahiert Text aus PDF"""
        if not self.pdf_available:
            raise RuntimeError("PyPDF2 not installed")
        
        import PyPDF2
        
        logger.info(f"Parsing PDF: {path.name}")
        
        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # Metadata
                metadata = {}
                if reader.metadata:
                    metadata = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'creator': reader.metadata.get('/Creator', ''),
                        'producer': reader.metadata.get('/Producer', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                    }
                
                # Text-Extraktion
                text_parts = []
                for page_num, page in enumerate(reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
                
                full_text = "\n\n".join(text_parts)
                
                # Cleanup
                full_text = self._clean_text(full_text)
                
                return ParsedDocument(
                    filename=path.name,
                    file_type='pdf',
                    text=full_text,
                    word_count=len(full_text.split()),
                    char_count=len(full_text),
                    metadata=metadata
                )
        
        except Exception as e:
            logger.error(f"Error parsing PDF {path.name}: {e}")
            raise
    
    def _parse_docx(self, path: Path) -> ParsedDocument:
        """
        Extrahiert Text aus DOCX
        
        Zwei Methoden:
        1. Pandoc (wenn verfügbar) - Beste Qualität
        2. python-docx (Fallback) - Standard
        """
        logger.info(f"Parsing DOCX: {path.name}")
        
        # Method 1: Pandoc (bevorzugt)
        if self.pandoc_available:
            return self._parse_docx_pandoc(path)
        
        # Method 2: python-docx (Fallback)
        if self.docx_available:
            return self._parse_docx_library(path)
        
        raise RuntimeError("No DOCX parser available. Install python-docx or pandoc")
    
    def _parse_docx_pandoc(self, path: Path) -> ParsedDocument:
        """
        DOCX-Parsing mit Pandoc (beste Qualität)
        
        Basierend auf DOCX Skill Documentation:
        - Behält Struktur bei (Überschriften, Listen, etc.)
        - Tracked Changes werden verarbeitet
        - Bessere Formatierung als python-docx
        """
        try:
            # Temporäre Markdown-Datei
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as tmp:
                tmp_path = tmp.name
            
            # Pandoc Konvertierung: DOCX → Markdown
            # --track-changes=accept: Akzeptiere alle Änderungen
            result = subprocess.run(
                ['pandoc', '--track-changes=accept', str(path), '-o', tmp_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                logger.warning(f"Pandoc error: {result.stderr}")
                raise subprocess.CalledProcessError(result.returncode, result.args)
            
            # Markdown lesen
            with open(tmp_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Cleanup temp file
            Path(tmp_path).unlink()
            
            # Markdown-Formatierung entfernen (optional)
            text = self._clean_markdown(text)
            text = self._clean_text(text)
            
            # Metadata (begrenzt bei Pandoc)
            metadata = {
                'parser': 'pandoc',
                'method': 'markdown_conversion'
            }
            
            return ParsedDocument(
                filename=path.name,
                file_type='docx',
                text=text,
                word_count=len(text.split()),
                char_count=len(text),
                metadata=metadata
            )
        
        except Exception as e:
            logger.warning(f"Pandoc parsing failed: {e}. Falling back to python-docx")
            if self.docx_available:
                return self._parse_docx_library(path)
            raise
    
    def _parse_docx_library(self, path: Path) -> ParsedDocument:
        """
        DOCX-Parsing mit python-docx (Fallback)
        
        Einfache Text-Extraktion ohne Formatierung
        """
        if not self.docx_available:
            raise RuntimeError("python-docx not installed")
        
        import docx
        
        try:
            doc = docx.Document(path)
            
            # Metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'parser': 'python-docx'
            }
            
            # Text-Extraktion (Paragraphs + Tables)
            text_parts = []
            
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            full_text = self._clean_text(full_text)
            
            return ParsedDocument(
                filename=path.name,
                file_type='docx',
                text=full_text,
                word_count=len(full_text.split()),
                char_count=len(full_text),
                metadata=metadata
            )
        
        except Exception as e:
            logger.error(f"Error parsing DOCX {path.name}: {e}")
            raise
    
    def _clean_markdown(self, text: str) -> str:
        """Entfernt Markdown-Formatierung"""
        # Headers (# ## ###)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        
        # Bold/Italic (**text** *text*)
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        
        # Links ([text](url))
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # Lists (- * +)
        text = re.sub(r'^[\-\*\+]\s+', '', text, flags=re.MULTILINE)
        
        return text
    
    def _clean_text(self, text: str) -> str:
        """Allgemeine Text-Bereinigung"""
        # Multiple whitespaces
        text = re.sub(r'\s+', ' ', text)
        
        # Multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Trim
        text = text.strip()
        
        return text
    
    def batch_parse(self, directory: str, 
                    pattern: str = "*",
                    recursive: bool = True) -> List[ParsedDocument]:
        """
        Parst mehrere Dokumente in einem Verzeichnis
        
        Args:
            directory: Verzeichnis-Pfad
            pattern: Datei-Pattern (z.B. "*.pdf", "*.docx", "*")
            recursive: Unterverzeichnisse durchsuchen
        
        Returns:
            Liste von ParsedDocument Objekten
        """
        dir_path = Path(directory)
        
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        # Sammle alle Dateien
        if recursive:
            files = list(dir_path.rglob(pattern))
        else:
            files = list(dir_path.glob(pattern))
        
        # Filter: nur .pdf und .docx
        supported_files = [
            f for f in files 
            if f.suffix.lower() in ['.pdf', '.docx', '.doc']
        ]
        
        logger.info(f"Found {len(supported_files)} documents to parse")
        
        # Parse alle Dateien
        results = []
        errors = []
        
        for file_path in supported_files:
            try:
                doc = self.parse_file(str(file_path))
                results.append(doc)
                logger.info(f"✓ Parsed: {file_path.name} ({doc.word_count} words)")
            except Exception as e:
                errors.append((file_path.name, str(e)))
                logger.error(f"✗ Failed: {file_path.name} - {e}")
        
        # Summary
        logger.info(f"\n{'='*60}")
        logger.info(f"BATCH PARSING COMPLETE")
        logger.info(f"{'='*60}")
        logger.info(f"Success: {len(results)}/{len(supported_files)}")
        logger.info(f"Failed: {len(errors)}")
        
        if errors:
            logger.warning("\nFailed files:")
            for filename, error in errors:
                logger.warning(f"  - {filename}: {error}")
        
        return results
    
    def get_statistics(self, documents: List[ParsedDocument]) -> Dict:
        """Berechnet Statistiken über geparste Dokumente"""
        if not documents:
            return {
                'total': 0,
                'by_type': {},
                'avg_words': 0,
                'total_words': 0
            }
        
        by_type = {}
        total_words = 0
        
        for doc in documents:
            by_type[doc.file_type] = by_type.get(doc.file_type, 0) + 1
            total_words += doc.word_count
        
        return {
            'total': len(documents),
            'by_type': by_type,
            'avg_words': total_words // len(documents),
            'total_words': total_words,
            'avg_chars': sum(d.char_count for d in documents) // len(documents)
        }


# ============================================================================
# EXAMPLE USAGE & TESTING
# ============================================================================

def test_parser():
    """Test-Funktion für Multi-Format Parser"""
    
    print(f"\n{'='*60}")
    print("MULTI-FORMAT DOCUMENT PARSER - TESTING")
    print(f"{'='*60}\n")
    
    # Initialize Parser
    parser = MultiFormatParser()
    
    # Test 1: PDF Check (wenn vorhanden)
    print("1. Checking PDF support...")
    if parser.pdf_available:
        print("   ✓ PyPDF2 available")
    else:
        print("   ✗ PyPDF2 not available")
        print("   Install: pip install PyPDF2 --break-system-packages")
    
    # Test 2: DOCX Check
    print("\n2. Checking DOCX support...")
    if parser.pandoc_available:
        print("   ✓ Pandoc available (preferred)")
    elif parser.docx_available:
        print("   ✓ python-docx available (fallback)")
    else:
        print("   ✗ No DOCX parser available")
        print("   Install: pip install python-docx --break-system-packages")
        print("   Or: sudo apt-get install pandoc")
    
    # Test 3: Beispiel-Text erstellen
    print("\n3. Creating test documents...")
    
    # Erstelle Test-DOCX
    try:
        import docx
        
        test_docx_path = "/home/claude/test_job_ad.docx"
        doc = docx.Document()
        doc.add_heading('UX Designer (m/w/d)', 0)
        doc.add_paragraph('''
        Wir suchen einen erfahrenen UX Designer mit folgenden Qualifikationen:
        
        Anforderungen:
        - 3+ Jahre Erfahrung im UX/UI Design
        - Expertise in Figma, Sketch, Adobe XD
        - Kenntnisse in User Research und Usability Testing
        - Erfahrung mit Agile/Scrum
        - Deutsch und Englisch fließend
        
        Nice to have:
        - HTML, CSS, JavaScript Grundkenntnisse
        - Design Systems Erfahrung
        - Prototyping Tools (InVision, Axure)
        ''')
        doc.save(test_docx_path)
        print(f"   ✓ Created test DOCX: {test_docx_path}")
        
        # Parse test DOCX
        print("\n4. Testing DOCX parsing...")
        parsed_docx = parser.parse_file(test_docx_path)
        print(f"   ✓ Parsed successfully!")
        print(f"   - File: {parsed_docx.filename}")
        print(f"   - Type: {parsed_docx.file_type}")
        print(f"   - Words: {parsed_docx.word_count}")
        print(f"   - Chars: {parsed_docx.char_count}")
        print(f"   - Parser: {parsed_docx.metadata.get('parser', 'N/A')}")
        print(f"\n   Text Preview (first 200 chars):")
        print(f"   {parsed_docx.text[:200]}...")
        
    except ImportError:
        print("   ✗ python-docx not available - skipping DOCX test")
    except Exception as e:
        print(f"   ✗ Error creating/parsing test DOCX: {e}")
    
    print(f"\n{'='*60}")
    print("✓ Testing complete!")
    print(f"{'='*60}\n")


def test_batch_parsing():
    """Test Batch-Parsing"""
    
    print(f"\n{'='*60}")
    print("BATCH PARSING TEST")
    print(f"{'='*60}\n")
    
    parser = MultiFormatParser()
    
    # Test mit Projekt-Dateien (falls vorhanden)
    test_dir = "/mnt/user-data/uploads"
    
    try:
        documents = parser.batch_parse(test_dir, pattern="*", recursive=False)
        
        if documents:
            stats = parser.get_statistics(documents)
            
            print(f"\n{'='*60}")
            print("STATISTICS")
            print(f"{'='*60}")
            print(f"Total documents: {stats['total']}")
            print(f"By type: {stats['by_type']}")
            print(f"Average words: {stats['avg_words']}")
            print(f"Total words: {stats['total_words']:,}")
        else:
            print("No documents found in test directory")
    
    except FileNotFoundError:
        print(f"Test directory not found: {test_dir}")
    except Exception as e:
        print(f"Error during batch parsing: {e}")


if __name__ == "__main__":
    print("="*60)
    print("MULTI-FORMAT DOCUMENT PARSER v2.0")
    print("Supports: PDF, DOCX, Google Docs (as DOCX)")
    print("="*60)
    
    test_parser()
    
    print("\n" + "="*60)
    print("For batch parsing, use:")
    print("  parser = MultiFormatParser()")
    print("  docs = parser.batch_parse('path/to/directory')")
    print("="*60)

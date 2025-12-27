from __future__ import annotations
import argparse, hashlib
from pathlib import Path
from datetime import datetime
from typing import List

from models.job_ad import JobAd
from services.preparation.data_preparation import DataPreparationService
from services.enrichment.organization_service import OrganizationEnrichmentService
from services.enrichment.role_service import RoleCategorizer
from services.extraction.competence_library import CompetenceLibrary
from services.extraction.competence_matcher import CompetenceMatcher
from services.analysis.analysis import QualityEvaluator
from services.reporting.reporting import ReportingService

try:
    import PyPDF2
    PDF_AVAILABLE = True
except Exception:
    PDF_AVAILABLE = False

try:
    from docx import Document
    WORD_AVAILABLE = True
except Exception:
    WORD_AVAILABLE = False

def read_pdf(path: Path) -> str:
    if not PDF_AVAILABLE: return ""
    try:
        text = ""
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for p in reader.pages:
                t = p.extract_text() or ""
                if t: text += t + "\n"
        return text.strip()
    except Exception:
        return ""

def read_docx(path: Path) -> str:
    if not WORD_AVAILABLE: return ""
    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""

def read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""

def collect_files(input_dir: Path) -> List[JobAd]:
    files = []
    for pat in ("*.pdf","*.docx","*.doc","*.txt"):
        files.extend(list(input_dir.rglob(pat)))
    jobs: List[JobAd] = []
    for fp in files:
        if fp.suffix.lower()==".pdf":
            raw = read_pdf(fp)
        elif fp.suffix.lower() in (".docx",".doc"):
            raw = read_docx(fp)
        else:
            raw = read_txt(fp)
        if not raw or len(raw) < 80:
            continue
        uid = hashlib.md5((fp.name + raw[:100]).encode("utf-8")).hexdigest()[:12]
        jobs.append(JobAd(
            id=uid,
            file_name=fp.name,
            source="local",
            raw_text=raw,
            char_count=len(raw),
            word_count=len(raw.split()),
            collection_date=datetime.now()
        ))
    return jobs

def main():
    ap = argparse.ArgumentParser(description="Job Mining – Modular Pipeline (stepwise learning)")
    ap.add_argument("--input", required=True, help="Input folder with job ads (.pdf/.docx/.txt)")
    ap.add_argument("--outdir", required=True, help="Output folder for CSV/JSON")
    ap.add_argument("--competence_csv", default="", help="Primary competence CSV (optional)")
    ap.add_argument("--competence_csv_extra", nargs="*", default=[], help="Extra competence CSVs (optional)")
    ap.add_argument("--no-spacy", action="store_true", help="Disable spaCy stage")
    ap.add_argument("--no-embeddings", action="store_true", help="Disable embeddings stage")
    args = ap.parse_args()

    input_dir = Path(args.input)
    outdir = Path(args.outdir); outdir.mkdir(parents=True, exist_ok=True)

    jobs = collect_files(input_dir)
    if not jobs:
        print("No valid files found.")
        return

    jobs = DataPreparationService().prepare(jobs)

    org = OrganizationEnrichmentService()
    role = RoleCategorizer()
    for j in jobs:
        j = org.enrich(j)
        j = role.categorize(j)

    lib = CompetenceLibrary(
        Path(args.competence_csv) if args.competence_csv else None,
        args.competence_csv_extra or None
    )
    matcher = CompetenceMatcher(lib, use_spacy=not args.no_spacy, use_embeddings=not args.no_embeddings)
    for j in jobs:
        text = j.cleaned_text or j.raw_text or ""
        j.competences = matcher.find(text)
        j.competences_count = len(j.competences)
        j.extraction_quality = (
            (1.0 if j.competences else 0.0) * 0.5 +
            (1.0 if (j.organization and j.organization.name and j.organization.name != "Unbekannt") else 0.0) * 0.25 +
            (1.0 if (j.location and j.location != "Nicht angegeben") else 0.0) * 0.25
        )

    metrics = QualityEvaluator().evaluate(jobs)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    rep = ReportingService()
    rep.export_csv(jobs, outdir / f"job_ads_{ts}.csv")
    rep.export_json(jobs, metrics, outdir / f"job_mining_results_{ts}.json")

    print("Done. Outputs at:", outdir)

if __name__ == "__main__":
    main()

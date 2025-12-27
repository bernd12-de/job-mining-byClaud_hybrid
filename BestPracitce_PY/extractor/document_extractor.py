"""
Domain Extractor: Document Extractor
Extrahiert Text aus PDF, DOCX, TXT, MD
"""

from pathlib import Path
from typing import Tuple, Optional
import re


class DocumentExtractor:
    """
    Extrahiert Text aus verschiedenen Dokumenten
    
    Unterstützt:
    - PDF (mit pdfplumber, fallback: PyPDF2)
    - DOCX (mit python-docx)
    - TXT, MD (native)
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']
    
    def extract(self, file_path: Path) -> Tuple[str, str]:
        """
        Extrahiere Text aus Datei
        
        Returns:
            (text, file_type)
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Datei nicht gefunden: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"Format nicht unterstützt: {suffix}")
        
        # Dispatch
        if suffix == '.pdf':
            text = self._extract_pdf(file_path)
        elif suffix == '.docx':
            text = self._extract_docx(file_path)
        elif suffix in ['.txt', '.md']:
            text = self._extract_text(file_path)
        else:
            raise ValueError(f"Unbekanntes Format: {suffix}")
        
        return text, suffix[1:]  # Ohne Punkt
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extrahiere aus PDF"""
        try:
            import pdfplumber
            
            text = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            return '\n\n'.join(text)
        
        except ImportError:
            # Fallback: PyPDF2
            try:
                import PyPDF2
                
                text = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                
                return '\n\n'.join(text)
            
            except ImportError:
                raise ImportError(
                    "Kein PDF-Parser verfügbar! "
                    "Installiere: pip install pdfplumber ODER pip install PyPDF2"
                )
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extrahiere aus DOCX"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Text aus Paragraphs
            text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Text aus Tabellen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n\n'.join(text)
        
        except ImportError:
            raise ImportError(
                "python-docx nicht verfügbar! "
                "Installiere: pip install python-docx"
            )
    
    def _extract_text(self, file_path: Path) -> str:
        """Extrahiere aus TXT/MD"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def clean_text(self, text: str) -> str:
        """
        Säubere extrahierten Text
        
        - Entferne Sonderzeichen
        - Normalisiere Whitespace
        - Entferne URLs
        """
        # Entferne URLs
        text = re.sub(r'https?://\S+', '', text)
        
        # Entferne Email
        text = re.sub(r'\S+@\S+', '', text)
        
        # Normalisiere Whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Entferne führende/trailing Whitespace
        text = text.strip()
        
        return text
    
    def get_preview(self, text: str, max_length: int = 200) -> str:
        """Erstelle Text-Preview"""
        if len(text) <= max_length:
            return text
        
        return text[:max_length] + '...'
